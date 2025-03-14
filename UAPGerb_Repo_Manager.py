import os
import markdown
from docx import Document

def convert_markdown_to_word(obsidian_folder, output_docx):
    if not os.path.exists(obsidian_folder):
        os.makedirs(obsidian_folder)
    
    doc = Document()
    doc.add_heading('Full Obsidian UAP Gerb Corpus', level=1)
    
    for root, _, files in os.walk(obsidian_folder):
        for filename in files:
            if filename.endswith(".md"):
                filepath = os.path.join(root, filename)
                print(f"Processing file: {filepath}")  # Debugging
                try:
                    with open(filepath, 'r', encoding='utf-8') as file:
                        md_content = file.read()
                        if not md_content.strip():
                            print(f"Warning: {filename} is empty.")
                            continue
                        
                        relative_path = os.path.relpath(filepath, obsidian_folder)
                        doc.add_heading(relative_path.replace('.md', ''), level=2)
                        doc.add_paragraph(md_content)
                        doc.add_page_break()
                except Exception as e:
                    print(f"Error reading {filename}: {e}")
    
    doc.save(output_docx)
    print(f"Conversion complete: {output_docx}")

def generate_directory_listing(obsidian_folder, output_txt):
    if not os.path.exists(os.path.dirname(output_txt)):
        os.makedirs(os.path.dirname(output_txt))

    with open(output_txt, 'w', encoding='utf-8') as file:
        file.write("Directory Listing for Obsidian Folder:\n\n")
        for root, dirs, files in os.walk(obsidian_folder):
            file.write(f"{root}:\n")
            for filename in files:
                file.write(f"  - {filename}\n")
    print(f"Directory listing saved: {output_txt}")

def generate_folder_tree(output_txt):
    with open(output_txt, 'w', encoding='utf-8') as file:
        for root, dirs, files in os.walk(os.getcwd()):
            level = root.replace(os.getcwd(), '').count(os.sep)
            indent = '    ' * level
            file.write(f"{indent}{os.path.basename(root)}/\n")
            subindent = '    ' * (level + 1)
            for filename in files:
                file.write(f"{subindent}{filename}\n")
    print(f"Folder tree saved: {output_txt}")

def convert_gerb_videos_to_word(obsidian_folder, output_folder, num_docs=5):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    gerb_videos_folder = os.path.join(obsidian_folder, "07 Gerbs UAP Videos")
    all_files = [f for f in os.listdir(gerb_videos_folder) if f.endswith(".md")]
    chunk_size = len(all_files) // num_docs + (len(all_files) % num_docs > 0)

    for i in range(num_docs):
        doc = Document()
        doc.add_heading(f'Gerb UAP Videos Part {i + 1}', level=1)
        start = i * chunk_size
        end = min((i + 1) * chunk_size, len(all_files))
        for filename in all_files[start:end]:
            filepath = os.path.join(gerb_videos_folder, filename)
            with open(filepath, 'r', encoding='utf-8') as file:
                md_content = file.read()
                doc.add_heading(filename.replace('.md', ''), level=2)
                doc.add_paragraph(md_content)
                doc.add_page_break()
        output_path = os.path.join(output_folder, f'Gerb_UAP_Videos_Part_{i + 1}.docx')
        doc.save(output_path)
        print(f"Saved {output_path}")

# Example usage
obsidian_folder = os.path.join(os.getcwd(), "UAPGerb")
output_folder = os.path.join(os.getcwd(), "AIGerbCorpus")
convert_markdown_to_word(obsidian_folder, os.path.join(output_folder, "Full UAP Gerb Video Obsidian Corpus.docx"))
#generate_directory_listing(obsidian_folder, os.path.join(output_folder, "directory_listing.txt"))
generate_folder_tree(os.path.join(output_folder, "folder_tree.txt"))
convert_gerb_videos_to_word(obsidian_folder, output_folder)
